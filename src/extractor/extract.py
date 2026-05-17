"""
extract.py — format-agnostic contract metadata extractor.
Single responsibility: bytes → ContractMetadata. No database, no storage calls.
"""

from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

# Add these specific imports for the gdoc handler
from google.auth import default as google_auth_default
from google.auth.exceptions import DefaultCredentialsError
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload


@dataclass
class ContractMetadata:
    source_file: str
    file_format: str
    client_name: str | None
    client_location: str | None
    provider_name: str | None
    provider_location: str | None
    effective_date: str | None
    expiration_date: str | None
    total_contract_value: float | None
    currency: str | None
    force_majeure_notice_days: int | None
    non_renewal_notice_months: int | None
    governing_law: str | None
    venue: str | None


# ── Format routers ─────────────────────────────────────────────────────────────


def _text_from_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _text_from_pdf(path: Path) -> str:
    import pdfplumber

    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if not text.strip():
                text = _ocr_fallback(page)
            pages.append(text)
    return _fix_ocr_noise("\n".join(pages))


def _ocr_fallback(page) -> str:
    try:
        import pytesseract

        return pytesseract.image_to_string(page.to_image(resolution=300).original, lang="eng")
    except Exception:
        return ""


def _fix_ocr_noise(text: str) -> str:
    text = re.sub(r"-\n(\w)", r"\1", text)
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"(?<!\w)€(?!\w)", "EUR ", text)
    return text


def _text_from_gdoc(doc_id: str) -> str:
    """Read a Google Drive file as text.

    Native Google Docs are exported as text. Uploaded .docx/PDF files are
    downloaded and parsed with the same local readers used for filesystem input.
    """
    try:
        sa_path = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
        scopes = ["https://www.googleapis.com/auth/drive.readonly"]

        if sa_path and os.path.exists(sa_path):
            creds = service_account.Credentials.from_service_account_file(sa_path, scopes=scopes)
        else:
            # This is where DefaultCredentialsError is usually raised
            creds, _ = google_auth_default(scopes=scopes)

        service = build("drive", "v3", credentials=creds)
        file_meta = (
            service.files()
            .get(fileId=doc_id, fields="id,name,mimeType", supportsAllDrives=True)
            .execute()
        )
        mime_type = file_meta.get("mimeType", "")
        name = file_meta.get("name", doc_id)

        if mime_type == "application/vnd.google-apps.folder":
            raise ValueError(
                f"Google Drive ID {doc_id} is a folder. Pass a file ID or use the pipeline CLI "
                "folder expansion."
            )

        if mime_type == "application/vnd.google-apps.document":
            request = service.files().export_media(fileId=doc_id, mimeType="text/plain")
            return _download_drive_media(request).decode("utf-8", errors="replace")

        raw_bytes = _download_drive_media(
            service.files().get_media(fileId=doc_id, supportsAllDrives=True)
        )
        return _text_from_drive_bytes(raw_bytes, name=name, mime_type=mime_type)

    except (DefaultCredentialsError, Exception) as e:
        raise OSError(f"Google Drive extraction failed: {e}") from e


def _download_drive_media(request) -> bytes:
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return buf.getvalue()


def _text_from_drive_bytes(raw_bytes: bytes, *, name: str, mime_type: str) -> str:
    suffix = Path(name).suffix.lower()
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        suffix = ".docx"
    elif mime_type == "application/pdf":
        suffix = ".pdf"

    if suffix == ".docx":
        from docx import Document

        doc = Document(io.BytesIO(raw_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if suffix == ".pdf":
        import pdfplumber

        pages = []
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return _fix_ocr_noise("\n".join(pages))

    raise ValueError(f"Unsupported Google Drive file type: {mime_type or suffix or 'unknown'}")


_READERS = {".docx": _text_from_docx, ".pdf": _text_from_pdf}
_GDOC_SCHEME = "gdoc://"


def _raw_text(raw_input: str) -> tuple[str, str]:
    if raw_input.startswith(_GDOC_SCHEME):
        doc_id = raw_input[len(_GDOC_SCHEME) :]
        return _text_from_gdoc(doc_id), "gdoc"
    path = Path(raw_input)
    suffix = path.suffix.lower()
    if suffix not in _READERS:
        raise ValueError(f"Unsupported format: {suffix!r}")
    return _READERS[suffix](path), suffix.lstrip(".")


# ── Field parsers ─────────────────────────────────────────────────────────────


def _parties(text: str) -> dict:
    between = re.search(
        r"BETWEEN[:\s]*\n+\**([^\*\n,]+)\**.*?located at\s+(.+?)\s*\(the\s+[\"“]?Client",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    and_ = re.search(
        r"\bAND[:\s]*\n+\**([^\*\n,]+)\**.*?located at\s+(.+?)\s*\(the\s+[\"“]?Provider",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    return {
        "client_name": between.group(1).strip() if between else None,
        "client_location": _clean_location(between.group(2)) if between else None,
        "provider_name": and_.group(1).strip() if and_ else None,
        "provider_location": _clean_location(and_.group(2)) if and_ else None,
    }


def _clean_location(raw: str) -> str:
    return re.sub(r"\s+", " ", raw).strip().rstrip(".,")


def _financial(text: str) -> dict:
    SYMBOLS = {"€": "EUR", "$": "USD", "£": "GBP"}

    def _to_float(raw: str) -> float:
        raw = raw.strip()
        if re.match(r"^\d{1,3}(\.\d{3})+(,\d+)?$", raw):
            raw = raw.replace(".", "").replace(",", ".")
        else:
            raw = raw.replace(",", "")
        return float(raw[:-1]) * 1000 if raw.lower().endswith("k") else float(raw)

    fee_block = re.search(
        r"(fixed fee|total.*?fee)(.*?)(term and expiration|\Z)",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    search_in = fee_block.group(2) if fee_block else text
    pattern = re.compile(
        r"([A-Z]{3}|€|\$|£)\s*([\d,\.]+k?)" r"|" r"([\d,\.]+k?)\s*([A-Z]{3})",
        re.IGNORECASE,
    )
    best_val, best_cur = None, None
    for m in pattern.finditer(search_in):
        raw_cur, raw_num = (m.group(1), m.group(2)) if m.group(1) else (m.group(4), m.group(3))
        try:
            val = _to_float(raw_num)
            cur = SYMBOLS.get(raw_cur.upper(), raw_cur.upper())
            if best_val is None or val > best_val:
                best_val, best_cur = val, cur
        except (ValueError, TypeError):
            continue
    return {"total_contract_value": best_val, "currency": best_cur}


def _dates(text: str) -> dict:
    DATE_PATS = [
        r"\b(\d{1,2}(?:st|nd|rd|th)?\s+of\s+\w+\s+\d{4})\b",
        r"\b(\w+\s+\d{1,2},\s+\d{4})\b",
    ]
    found = []
    for pat in DATE_PATS:
        for m in re.finditer(pat, text, re.IGNORECASE):
            raw = re.sub(r"(st|nd|rd|th)\s+of\s+", " ", m.group(1), flags=re.IGNORECASE)
            raw = re.sub(r"(st|nd|rd|th)\b", "", raw, flags=re.IGNORECASE).strip()
            for fmt in ("%B %d, %Y", "%d %B %Y"):
                try:
                    found.append(datetime.strptime(raw, fmt))
                    break
                except ValueError:
                    pass

    eff = min(found).strftime("%Y-%m-%d") if found else None
    exp = max(found).strftime("%Y-%m-%d") if found else None

    if re.search(r"expire.*?after one year", text, re.IGNORECASE) and eff:
        d = datetime.strptime(eff, "%Y-%m-%d")
        try:
            exp = d.replace(year=d.year + 1).strftime("%Y-%m-%d")
        except ValueError:
            exp = d.replace(year=d.year + 1, day=28).strftime("%Y-%m-%d")

    return {"effective_date": eff, "expiration_date": exp}


def _obligations(text: str) -> dict:
    fm = re.search(
        r"force majeure.*?exceeding\s+(?:\w+\s+)?\(?\s*(\d+)\s*\)?\s*consecutive\s*days",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    nr = re.search(
        r"non.renewal.*?(?:at least\s+)?(?:\w+\s+)?\(?\s*(\d+)\s*\)?\s*months?",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    return {
        "force_majeure_notice_days": int(fm.group(1)) if fm else None,
        "non_renewal_notice_months": int(nr.group(1)) if nr else None,
    }


def _law(text: str) -> dict:
    law = re.search(
        r"governing law.*?laws? of (?:the\s+)?([^\.\n]+)", text, re.IGNORECASE | re.DOTALL
    )
    venue = re.search(
        r"(?:venue|place of jurisdiction).*?(?:shall be|is)\s+([A-Z][a-zA-Z\s]+)",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    return {
        "governing_law": law.group(1).strip() if law else None,
        "venue": venue.group(1).strip() if venue else None,
    }


# ── Public API ─────────────────────────────────────────────────────────────────


def extract_contract_metadata(path: str | Path) -> ContractMetadata:
    raw = str(path)
    if raw.startswith(_GDOC_SCHEME):
        source_name = raw
    else:
        source_name = Path(path).name
    text, fmt = _raw_text(raw)
    return ContractMetadata(
        source_file=source_name,
        file_format=fmt,
        **_parties(text),
        **_financial(text),
        **_dates(text),
        **_obligations(text),
        **_law(text),
    )
